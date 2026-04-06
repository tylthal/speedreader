"""DOCX parser: extracts metadata and chapters from .docx files."""

from __future__ import annotations

import re
import uuid

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from backend.parser_types import InlineImage, ParsedBook, ParsedChapter, save_inline_image

_MIN_CHAPTER_LENGTH = 50
_FALLBACK_WORD_LIMIT = 3000
_MIN_IMAGE_SIZE = 500  # bytes — skip tiny images


def _extract_paragraph_images(
    para,
    doc: Document,
    pub_uuid: str,
    counter: int,
) -> tuple[int, list[InlineImage]]:
    """Extract images from a paragraph's inline shapes.

    Returns (next_counter, list_of_inline_images).
    """
    inline_images: list[InlineImage] = []

    # Access images through the paragraph's XML — look for blip elements
    # which reference image relationships
    ns = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }

    for blip in para._element.findall(".//a:blip", ns):
        embed_id = blip.get(f"{{{ns['r']}}}embed")
        if not embed_id:
            continue

        try:
            rel = para.part.rels.get(embed_id)
            if rel is None:
                continue
            image_part = rel.target_part
            image_data = image_part.blob
        except Exception:
            continue

        if not image_data or len(image_data) < _MIN_IMAGE_SIZE:
            continue

        # Determine MIME type from content type
        content_type = getattr(image_part, "content_type", "image/jpeg")
        placeholder = f"{{{{IMG_{counter}}}}}"
        rel_path, width, height = save_inline_image(
            image_data, pub_uuid, counter, content_type,
        )

        inline_images.append(InlineImage(
            placeholder=placeholder,
            image_path=rel_path,
            alt="",
            width=width,
            height=height,
            mime_type=content_type,
        ))
        counter += 1

    return counter, inline_images


def parse_docx(file_path: str) -> ParsedBook:
    """Parse a DOCX file into chapters with inline images."""
    doc = Document(file_path)
    pub_uuid = str(uuid.uuid4())
    image_counter = 0

    # Metadata
    title = (doc.core_properties.title or "").strip() or "Untitled"
    author = (doc.core_properties.author or "").strip() or "Unknown Author"

    # Group paragraphs by heading boundaries
    chapters: list[ParsedChapter] = []
    current_title = ""
    current_parts: list[str] = []
    current_images: list[InlineImage] = []

    for para in doc.paragraphs:
        style_name = (para.style.name or "").lower() if para.style else ""
        text = para.text.strip()

        # Extract images from this paragraph
        image_counter, para_images = _extract_paragraph_images(
            para, doc, pub_uuid, image_counter,
        )

        # Check if this is a heading paragraph
        is_heading = (
            style_name.startswith("heading")
            or style_name.startswith("titre")
        )

        if is_heading:
            # Save previous chapter
            if current_parts or current_images:
                body = "\n".join(current_parts).strip()
                if len(body) >= _MIN_CHAPTER_LENGTH or current_images:
                    ch_title = current_title or f"Chapter {len(chapters) + 1}"
                    chapters.append(ParsedChapter(
                        title=ch_title, text=body, inline_images=current_images,
                    ))

            current_title = text
            current_parts = []
            current_images = []
        else:
            if para_images:
                # Insert placeholders into the text
                placeholders = " ".join(img.placeholder for img in para_images)
                if text:
                    text = f"{placeholders} {text}"
                else:
                    text = placeholders
                current_images.extend(para_images)
            if text:
                current_parts.append(text)

    # Final chapter
    if current_parts or current_images:
        body = "\n".join(current_parts).strip()
        if len(body) >= _MIN_CHAPTER_LENGTH or current_images:
            ch_title = current_title or f"Chapter {len(chapters) + 1}"
            chapters.append(ParsedChapter(
                title=ch_title, text=body, inline_images=current_images,
            ))

    # Fallback if no headings found
    if not chapters:
        all_parts: list[str] = []
        all_images: list[InlineImage] = []
        img_ctr = 0
        for para in doc.paragraphs:
            text = para.text.strip()
            img_ctr, para_imgs = _extract_paragraph_images(para, doc, pub_uuid, img_ctr)
            if para_imgs:
                placeholders = " ".join(img.placeholder for img in para_imgs)
                text = f"{placeholders} {text}" if text else placeholders
                all_images.extend(para_imgs)
            if text:
                all_parts.append(text)

        all_text = "\n".join(all_parts)
        if len(all_text) >= _MIN_CHAPTER_LENGTH:
            words = all_text.split()
            for start in range(0, len(words), _FALLBACK_WORD_LIMIT):
                chunk = " ".join(words[start:start + _FALLBACK_WORD_LIMIT]).strip()
                if len(chunk) >= _MIN_CHAPTER_LENGTH:
                    chapters.append(ParsedChapter(
                        title=f"Chapter {len(chapters) + 1}",
                        text=chunk,
                        inline_images=all_images if start == 0 else [],
                    ))

    return ParsedBook(title=title, author=author, chapters=chapters)
